"""Unit tests for .ai/templates/docx_helper.py.

Builds a real document via the helper API (create_doc + add_* functions), saves
it to a temp .docx, then asserts: (1) harness_core.docx_validation accepts it as
valid OOXML, (2) word/document.xml contains a table and a non-empty body, and
(3) format invariants the helper actually sets (A4 page geometry, Arial latin /
맑은 고딕 east-asia run fonts, 12pt body, code-block shading + Courier New, fixed
table layout at 15.99cm with Table Grid style). Skipped entirely when
python-docx is not installed.
"""
from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from support import ensure_ai_path

_AI_DIR = ensure_ai_path()
_TEMPLATES_DIR = _AI_DIR / "templates"

from harness_core.docx_validation import validate_docx_file  # noqa: E402

try:
    import docx  # noqa: F401
    _HAS_PYTHON_DOCX = True
except ImportError:
    _HAS_PYTHON_DOCX = False

if _HAS_PYTHON_DOCX:
    if str(_TEMPLATES_DIR) not in sys.path:
        sys.path.insert(0, str(_TEMPLATES_DIR))
    import docx_helper  # noqa: E402
    from docx import Document  # noqa: E402
    from docx.oxml.ns import qn  # noqa: E402
    from docx.shared import Pt, RGBColor  # noqa: E402

BODY_TEXT = "Body paragraph text."
TITLE_TEXT = "stock-report specification"
CODE_TEXT = "print('hello')"


@unittest.skipUnless(_HAS_PYTHON_DOCX, "python-docx (docx) is not installed")
class DocxHelperTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        super().setUpClass()
        cls._tmp = tempfile.TemporaryDirectory()
        cls.addClassCleanup(cls._tmp.cleanup)
        cls.docx_path = Path(cls._tmp.name) / "spec.docx"

        doc = docx_helper.create_doc()
        docx_helper.add_document_title(doc, TITLE_TEXT)
        docx_helper.add_h1(doc, "1. Overview")
        docx_helper.add_h2(doc, "1.1 Scope")
        docx_helper.add_paragraph(doc, BODY_TEXT)
        docx_helper.add_bullet(doc, "First bullet item")
        docx_helper.add_table(
            doc,
            ["Name", "Value"],
            [["alpha", "1"], ["beta", "2"]],
        )
        docx_helper.add_code_block(doc, CODE_TEXT)
        docx_helper.add_source_note(doc, "Source: internal test fixture")
        doc.save(str(cls.docx_path))

        cls.reopened = Document(str(cls.docx_path))

    def _paragraph_with_text(self, wanted: str):
        for paragraph in self.reopened.paragraphs:
            if paragraph.text == wanted:
                return paragraph
        self.fail(f"No paragraph with text {wanted!r} in saved document")

    def test_validate_docx_file_accepts_output(self) -> None:
        self.assertIsNone(validate_docx_file(self.docx_path))

    def test_document_xml_has_table_and_non_empty_body(self) -> None:
        with zipfile.ZipFile(self.docx_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        self.assertTrue(xml.strip())
        self.assertIn("<w:tbl", xml)
        self.assertIn("<w:p", xml)
        self.assertIn(BODY_TEXT, xml)
        # East-asia font token is written into the run properties.
        self.assertIn("맑은 고딕", xml)

    def test_page_geometry_is_a4_with_template_margins(self) -> None:
        sec = self.reopened.sections[0]
        self.assertAlmostEqual(sec.page_width.cm, 21.0, delta=0.01)
        self.assertAlmostEqual(sec.page_height.cm, 29.7, delta=0.01)
        self.assertAlmostEqual(sec.top_margin.cm, 3.0, delta=0.01)
        self.assertAlmostEqual(sec.bottom_margin.cm, 2.54, delta=0.01)
        self.assertAlmostEqual(sec.left_margin.cm, 2.54, delta=0.01)
        self.assertAlmostEqual(sec.right_margin.cm, 2.54, delta=0.01)
        self.assertAlmostEqual(sec.header_distance.cm, 1.5, delta=0.01)
        self.assertAlmostEqual(sec.footer_distance.cm, 1.75, delta=0.01)

    def test_body_run_uses_arial_malgun_gothic_at_12pt(self) -> None:
        run = self._paragraph_with_text(BODY_TEXT).runs[0]
        self.assertEqual(run.font.name, "Arial")
        self.assertEqual(run.font.size, Pt(12))
        rfonts = run._element.rPr.rFonts
        self.assertEqual(rfonts.get(qn("w:ascii")), "Arial")
        self.assertEqual(rfonts.get(qn("w:hAnsi")), "Arial")
        self.assertEqual(rfonts.get(qn("w:eastAsia")), "맑은 고딕")
        self.assertEqual(rfonts.get(qn("w:cs")), "Arial")

    def test_title_run_is_16pt_bold_dark_blue(self) -> None:
        run = self._paragraph_with_text(TITLE_TEXT).runs[0]
        self.assertTrue(run.font.bold)
        self.assertEqual(run.font.size, Pt(16))
        self.assertEqual(run.font.color.rgb, RGBColor(0x00, 0x20, 0x60))

    def test_code_block_uses_courier_9pt_with_gray_shading(self) -> None:
        paragraph = self._paragraph_with_text(CODE_TEXT)
        run = paragraph.runs[0]
        self.assertEqual(run.font.name, "Courier New")
        self.assertEqual(run.font.size, Pt(9))
        shd = paragraph._p.pPr.find(qn("w:shd"))
        self.assertIsNotNone(shd)
        self.assertEqual(shd.get(qn("w:fill")), "F2F2F2")

    def test_table_grid_fixed_layout_and_total_width(self) -> None:
        table = self.reopened.tables[0]
        self.assertEqual(table.style.name, "Table Grid")
        self.assertEqual(table.rows[0].cells[0].text, "Name")
        self.assertEqual(table.rows[1].cells[1].text, "1")
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        self.assertIsNotNone(layout)
        self.assertEqual(layout.get(qn("w:type")), "fixed")
        tbl_w = tbl_pr.find(qn("w:tblW"))
        self.assertIsNotNone(tbl_w)
        self.assertEqual(tbl_w.get(qn("w:type")), "dxa")
        expected_twips = docx_helper._cm_to_twips(docx_helper.TABLE_TOTAL_WIDTH_CM)
        self.assertEqual(tbl_w.get(qn("w:w")), str(expected_twips))
        borders = tbl_pr.find(qn("w:tblBorders"))
        self.assertIsNotNone(borders)
        top = borders.find(qn("w:top"))
        self.assertEqual(top.get(qn("w:val")), "single")
        self.assertEqual(top.get(qn("w:sz")), docx_helper.TABLE_BORDER_SIZE)


if __name__ == "__main__":
    unittest.main()
