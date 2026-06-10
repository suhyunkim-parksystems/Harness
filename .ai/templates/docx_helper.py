"""
.ai/templates/docx_helper.py

명세서 .docx 생성용 python-docx 헬퍼.
06_document 단계에서 이 파일을 import하여 내용만 채운다.

사용법:
    import sys
    sys.path.insert(0, ".ai/templates")
    from docx_helper import (
        create_doc, add_document_kicker, add_document_title,
        add_h1, add_h2, add_paragraph, add_bullet,
        add_code_block, add_table
    )

    doc = create_doc()
    add_document_kicker(doc, "Feature Specification")
    add_document_title(doc, "feature-name 명세서")
    add_h1(doc, "1. 개요")
    add_paragraph(doc, "기능 목적 요약...")
    doc.save(".project/docs/feature-name_명세서.docx")
"""

from __future__ import annotations

from typing import Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# 문서 형식 토큰. 참조 문서의 OOXML에서 추출한 값만 둔다.
LATIN_FONT = "Arial"
EAST_ASIA_FONT = "맑은 고딕"
COMPLEX_SCRIPT_FONT = "Arial"

COLOR_BLACK = "000000"
COLOR_DARK_BLUE = "002060"
COLOR_ACCENT_BLUE = "00B0F0"
COLOR_MUTED_GRAY = "808080"
COLOR_CODE_BG = "F2F2F2"

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_CM = 3.0
MARGIN_RIGHT_CM = 2.54
MARGIN_BOTTOM_CM = 2.54
MARGIN_LEFT_CM = 2.54
HEADER_DISTANCE_CM = 1.5
FOOTER_DISTANCE_CM = 1.75

BODY_SIZE_PT = 12
SMALL_SIZE_PT = 10
TITLE_SIZE_PT = 16
SECTION_SIZE_PT = 16
SUBSECTION_SIZE_PT = 12
DEFAULT_LINE_SPACING = 1.5

TABLE_TOTAL_WIDTH_CM = 15.99
TABLE_BORDER_SIZE = "4"  # OOXML half-points: 4 = 0.5pt
TABLE_CELL_MARGIN_DXA = "108"  # 약 0.19cm
TABLE_ROW_HEIGHT_PT = 20.2


def create_doc() -> Document:
    """A4, 표준 문서 형식 토큰이 적용된 Document 객체를 반환한다."""
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(PAGE_WIDTH_CM)
    sec.page_height = Cm(PAGE_HEIGHT_CM)
    sec.left_margin = Cm(MARGIN_LEFT_CM)
    sec.right_margin = Cm(MARGIN_RIGHT_CM)
    sec.top_margin = Cm(MARGIN_TOP_CM)
    sec.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    sec.header_distance = Cm(HEADER_DISTANCE_CM)
    sec.footer_distance = Cm(FOOTER_DISTANCE_CM)

    _configure_doc_defaults(doc)
    _configure_paragraph_style(
        doc.styles["Normal"],
        size_pt=11,
        line_spacing=DEFAULT_LINE_SPACING,
        space_after_pt=8,
    )
    _configure_paragraph_style(
        doc.styles["Heading 1"],
        size_pt=SECTION_SIZE_PT,
        bold=True,
        color=COLOR_BLACK,
        line_spacing=DEFAULT_LINE_SPACING,
        space_before_pt=0,
        space_after_pt=0,
    )
    _configure_paragraph_style(
        doc.styles["Heading 2"],
        size_pt=SUBSECTION_SIZE_PT,
        bold=True,
        color=COLOR_BLACK,
        line_spacing=DEFAULT_LINE_SPACING,
        space_before_pt=0,
        space_after_pt=0,
    )
    _configure_paragraph_style(
        doc.styles["List Bullet"],
        size_pt=BODY_SIZE_PT,
        line_spacing=DEFAULT_LINE_SPACING,
        space_after_pt=0,
    )
    if "No Spacing" in [style.name for style in doc.styles]:
        _configure_paragraph_style(
            doc.styles["No Spacing"],
            size_pt=BODY_SIZE_PT,
            line_spacing=DEFAULT_LINE_SPACING,
            space_after_pt=0,
        )

    return doc


def add_document_kicker(doc: Document, text: str):
    """상단 보조 문구. 10pt 회색 텍스트."""
    p = doc.add_paragraph()
    _format_paragraph(p, line_spacing=DEFAULT_LINE_SPACING, space_after_pt=0)
    run = p.add_run(text)
    _format_run(run, size_pt=SMALL_SIZE_PT, color=COLOR_MUTED_GRAY)
    return p


def add_document_title(doc: Document, text: str, accent_text: str | None = None):
    """16pt bold 제목. accent_text가 있으면 해당 부분만 청록색으로 강조한다."""
    p = doc.add_paragraph()
    _format_paragraph(p, line_spacing=DEFAULT_LINE_SPACING, space_after_pt=0)
    if accent_text and accent_text in text:
        before, after = text.split(accent_text, 1)
        if before:
            _format_run(p.add_run(before), size_pt=TITLE_SIZE_PT, bold=True, color=COLOR_DARK_BLUE)
        _format_run(p.add_run(accent_text), size_pt=TITLE_SIZE_PT, bold=True, color=COLOR_ACCENT_BLUE)
        if after:
            _format_run(p.add_run(after), size_pt=TITLE_SIZE_PT, bold=True, color=COLOR_DARK_BLUE)
    else:
        _format_run(p.add_run(text), size_pt=TITLE_SIZE_PT, bold=True, color=COLOR_DARK_BLUE)
    return p


def add_h1(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    _format_paragraph(p, line_spacing=DEFAULT_LINE_SPACING, space_after_pt=0)
    for run in p.runs:
        _format_run(run, size_pt=SECTION_SIZE_PT, bold=True, color=COLOR_BLACK)
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    _format_paragraph(p, line_spacing=DEFAULT_LINE_SPACING, space_after_pt=0)
    for run in p.runs:
        _format_run(run, size_pt=SUBSECTION_SIZE_PT, bold=True, color=COLOR_BLACK)
    return p


def add_paragraph(doc: Document, text: str, *, bold: bool = False):
    p = doc.add_paragraph()
    _format_paragraph(p, line_spacing=DEFAULT_LINE_SPACING, space_after_pt=0)
    run = p.add_run(text)
    _format_run(run, size_pt=BODY_SIZE_PT, bold=bold, color=COLOR_BLACK)
    return p


def add_source_note(doc: Document, text: str):
    """출처/단위/각주성 문구. 10pt 본문."""
    p = doc.add_paragraph()
    _format_paragraph(p, line_spacing=DEFAULT_LINE_SPACING, space_after_pt=0)
    run = p.add_run(text)
    _format_run(run, size_pt=SMALL_SIZE_PT, color=COLOR_BLACK)
    return p


def add_unit_label(doc: Document, text: str):
    """표 위 단위 표기. 10pt, 오른쪽 정렬."""
    p = doc.add_paragraph()
    _format_paragraph(
        p,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        line_spacing=DEFAULT_LINE_SPACING,
        space_after_pt=0,
    )
    run = p.add_run(text)
    _format_run(run, size_pt=SMALL_SIZE_PT, color=COLOR_BLACK)
    return p


def add_bullet(doc: Document, text: str, *, bold: bool = False):
    """불릿 목록 한 줄 추가."""
    p = doc.add_paragraph(style="List Bullet")
    _format_paragraph(
        p,
        line_spacing=DEFAULT_LINE_SPACING,
        space_after_pt=0,
        left_indent_cm=0.5,
        first_line_indent_cm=-0.5,
    )
    run = p.add_run(text)
    _format_run(run, size_pt=BODY_SIZE_PT, bold=bold, color=COLOR_BLACK)
    return p


def add_code_block(doc: Document, text: str):
    """회색 배경 + Courier New 고정폭 코드 블록."""
    p = doc.add_paragraph()
    _format_paragraph(
        p,
        line_spacing=1.0,
        space_before_pt=4,
        space_after_pt=4,
        left_indent_cm=0.5,
    )

    pPr = p._p.get_or_add_pPr()
    shd = _get_or_add_child(pPr, "w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), COLOR_CODE_BG)

    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


def add_table(
    doc: Document,
    headers: Sequence[object],
    rows: Sequence[Sequence[object]],
    *,
    column_widths_cm: Sequence[float] | None = None,
    font_size_pt: int = SMALL_SIZE_PT,
    header_font_size_pt: int | None = None,
    alignments: Sequence[str] | None = None,
    row_height_pt: float = TABLE_ROW_HEIGHT_PT,
):
    """얇은 단일선 그리드 표를 추가한다.

    headers : ['컬럼1', '컬럼2', ...]
    rows    : [['값1', '값2', ...], ...]

    column_widths_cm을 생략하면 전체 폭 15.99cm를 균등 분배한다.
    alignments는 'left', 'center', 'right' 문자열 목록이다.
    """
    if not headers:
        raise ValueError("headers must contain at least one column.")

    column_count = len(headers)
    widths = _normalize_column_widths(column_count, column_widths_cm)
    header_font_size = header_font_size_pt or font_size_pt

    table = doc.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_width(table, TABLE_TOTAL_WIDTH_CM)
    _set_table_borders(table)
    _set_table_cell_margins(table)
    _set_table_fixed_layout(table)

    _fill_table_row(
        table.rows[0],
        headers,
        widths,
        font_size_pt=header_font_size,
        bold=True,
        alignments=alignments,
        row_height_pt=row_height_pt,
    )

    for row_data in rows:
        row_values = list(row_data)
        if len(row_values) < column_count:
            row_values.extend([""] * (column_count - len(row_values)))
        _fill_table_row(
            table.add_row(),
            row_values[:column_count],
            widths,
            font_size_pt=font_size_pt,
            bold=False,
            alignments=alignments,
            row_height_pt=row_height_pt,
        )

    doc.add_paragraph()
    return table


def _configure_doc_defaults(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)

    rPr_default = _get_or_add_child(doc_defaults, "w:rPrDefault", first=True)
    rPr = _get_or_add_child(rPr_default, "w:rPr")
    _set_rpr_fonts(rPr)
    _set_rpr_size(rPr, 11)
    _set_rpr_color(rPr, COLOR_BLACK)

    pPr_default = _get_or_add_child(doc_defaults, "w:pPrDefault")
    pPr = _get_or_add_child(pPr_default, "w:pPr")
    _set_ppr_spacing(pPr, line_spacing=DEFAULT_LINE_SPACING, space_after_pt=8)


def _configure_paragraph_style(
    style,
    *,
    size_pt: float,
    bold: bool = False,
    color: str = COLOR_BLACK,
    line_spacing: float | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
) -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = _rgb(color)
    _set_rpr_fonts(style._element.get_or_add_rPr())
    pf = style.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)


def _format_paragraph(
    paragraph,
    *,
    alignment=None,
    line_spacing: float | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
    left_indent_cm: float | None = None,
    first_line_indent_cm: float | None = None,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if space_before_pt is not None:
        fmt.space_before = Pt(space_before_pt)
    if space_after_pt is not None:
        fmt.space_after = Pt(space_after_pt)
    if left_indent_cm is not None:
        fmt.left_indent = Cm(left_indent_cm)
    if first_line_indent_cm is not None:
        fmt.first_line_indent = Cm(first_line_indent_cm)


def _format_run(
    run,
    *,
    size_pt: float,
    bold: bool = False,
    color: str = COLOR_BLACK,
) -> None:
    run.font.name = LATIN_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = _rgb(color)
    rPr = run._element.get_or_add_rPr()
    _set_rpr_fonts(rPr)
    _set_kern_zero(rPr)


def _set_rpr_fonts(rPr) -> None:
    rFonts = _get_or_add_child(rPr, "w:rFonts", first=True)
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rFonts.set(qn("w:cs"), COMPLEX_SCRIPT_FONT)
    rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def _set_rpr_size(rPr, size_pt: float) -> None:
    value = str(int(round(size_pt * 2)))
    for tag in ("w:sz", "w:szCs"):
        el = _get_or_add_child(rPr, tag)
        el.set(qn("w:val"), value)


def _set_rpr_color(rPr, color: str) -> None:
    el = _get_or_add_child(rPr, "w:color")
    el.set(qn("w:val"), color)


def _set_kern_zero(rPr) -> None:
    kern = _get_or_add_child(rPr, "w:kern")
    kern.set(qn("w:val"), "0")


def _set_ppr_spacing(
    pPr,
    *,
    line_spacing: float | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
) -> None:
    spacing = _get_or_add_child(pPr, "w:spacing")
    if line_spacing is not None:
        spacing.set(qn("w:line"), str(int(round(line_spacing * 240))))
        spacing.set(qn("w:lineRule"), "auto")
    if space_before_pt is not None:
        spacing.set(qn("w:before"), str(_pt_to_twips(space_before_pt)))
    if space_after_pt is not None:
        spacing.set(qn("w:after"), str(_pt_to_twips(space_after_pt)))


def _fill_table_row(
    row,
    values: Sequence[object],
    widths_cm: Sequence[float],
    *,
    font_size_pt: int,
    bold: bool,
    alignments: Sequence[str] | None,
    row_height_pt: float,
) -> None:
    row.height = Pt(row_height_pt)
    for index, value in enumerate(values):
        cell = row.cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_width(cell, widths_cm[index])
        cell.text = ""
        paragraph = cell.paragraphs[0]
        alignment = _column_alignment(index, alignments)
        _format_paragraph(paragraph, alignment=alignment, space_after_pt=0)
        run = paragraph.add_run(str(value))
        _format_run(run, size_pt=font_size_pt, bold=bold, color=COLOR_BLACK)


def _column_alignment(index: int, alignments: Sequence[str] | None):
    if alignments and index < len(alignments):
        value = str(alignments[index]).strip().lower()
        if value == "left":
            return WD_ALIGN_PARAGRAPH.LEFT
        if value == "right":
            return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.CENTER


def _normalize_column_widths(
    column_count: int, column_widths_cm: Sequence[float] | None
) -> list[float]:
    if column_widths_cm and len(column_widths_cm) == column_count:
        return [float(width) for width in column_widths_cm]
    equal = TABLE_TOTAL_WIDTH_CM / column_count
    return [equal] * column_count


def _set_table_width(table, width_cm: float) -> None:
    tblPr = table._tbl.tblPr
    tblW = _get_or_add_child(tblPr, "w:tblW", first=True)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(_cm_to_twips(width_cm)))


def _set_table_fixed_layout(table) -> None:
    tblPr = table._tbl.tblPr
    layout = _get_or_add_child(tblPr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")


def _set_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    borders = _get_or_add_child(tblPr, "w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = _get_or_add_child(borders, f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), TABLE_BORDER_SIZE)
        el.set(qn("w:color"), "auto")
        el.set(qn("w:space"), "0")


def _set_table_cell_margins(table) -> None:
    tblPr = table._tbl.tblPr
    margins = _get_or_add_child(tblPr, "w:tblCellMar")
    values = {
        "top": "0",
        "left": TABLE_CELL_MARGIN_DXA,
        "bottom": "0",
        "right": TABLE_CELL_MARGIN_DXA,
    }
    for side, width in values.items():
        el = _get_or_add_child(margins, f"w:{side}")
        el.set(qn("w:w"), width)
        el.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_cm: float) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = _get_or_add_child(tcPr, "w:tcW", first=True)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(_cm_to_twips(width_cm)))


def _get_or_add_child(parent, tag: str, *, first: bool = False):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        if first:
            parent.insert(0, child)
        else:
            parent.append(child)
    return child


def _cm_to_twips(value: float) -> int:
    return int(round(value / 2.54 * 1440))


def _pt_to_twips(value: float) -> int:
    return int(round(value * 20))


def _rgb(hex_color: str) -> RGBColor:
    value = hex_color.strip().lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))
